# importing package
from openpyxl import load_workbook
import os
import pandas as pd
from tqdm import tqdm
import sys
import time
import datetime
import xlsxwriter
from itertools import islice
from xlwt import Workbook
import docx

def get_column_headers(worksheet_name, sheet_name):
    """
    A function to get the column headers for an excel sheet.
    :param worksheet_name:
    :param sheet_name:
    :return:
    """
    datasets = pd.read_excel(worksheet_name, sheet_name=sheet_name)
    column_headers = datasets.columns
    column_headers = list(column_headers)
    return column_headers

def read_column(column_name, input_datasets):
    """
    A function to read a particular column values and return the parameters as list.
    :param column_name:
    :return:
    """
    column_values = input_datasets[column_name]
    column_values = column_values.dropna().tolist()
    return column_values

def eliminate_empty_lists(keywords_headers, keyword_excel_data, number):
    # Still a work in progress
    """
    A function to compare and eliminate a header with 0 values
    :param keywords_headers:
    :param keyword_excel_data:
    :return:
    """

    for values in keywords_headers :
        column_values = read_column(values, keyword_excel_data)
        if len(column_values) == 0:
            number = keywords_headers.index()
            keywords_headers.pop(number)

    return keywords_headers

def folder_name():
    """
    A function to return the name of the folder when the program runs
    :return:
    """
    # using now() to get current time
    current_time = datetime.datetime.now()

    # get the first 3 characters of month name
    monthinteger = int(current_time.month)
    month = datetime.date(1900, monthinteger, 1).strftime('%B')
    month = month[:3]

    # now we create the folder name
    folder_name = str(current_time.day) + "-" + str(month) + "-" + str(current_time.year) + "-" + str(
        current_time.hour) + "-" + str(current_time.minute) + "-" + str(current_time.second)

    return folder_name


def create_unique_excels(keywords_headers, keyword_excel_data, main_excel_data, main_column_name, number, folder_name, remed_folder, remed_excel_data):
    """
    A function to divide main-data using the keywords provided by the client
    :param keywords_headers:
    :param keyword_excel_data:
    :param main_excel_data:
    :param main_column_name:
    :param number:
    :param folder_name:
    :return:
    """

    # create the excel sheet name
    create_excel = str(keywords_headers[number]) + ".xlsx"
    create_doc = str(keywords_headers[number]) + ".docx"

    # create values for writing document
    heading = str(keywords_headers[number]) + " - Remediation Plan"
    list_to_write = remed_excel_data[keywords_headers[number]]

    # create the folder name
    file_path = str(folder_name) + "/" + str(create_excel)
    remed_filepath = str(remed_folder) + "/" + str(create_doc)
    appended_data = []

    # read the column values for each keyword provided
    column_values = read_column(keywords_headers[number], keyword_excel_data)
    # create a lowercase version of the keywords and add it to the main keywords list
    l_cv = [i.lower() for i in column_values]
    # create a uppercase version of the keywords and add it to the main keywords list
    u_cv = [i.upper() for i in column_values]
    # add the lowercase and uppercase values to the original list
    [column_values.extend(l) for l in (l_cv, u_cv)]

    # check if the traversed column values is not equal to zero
    if len(column_values) != 0 :

        # traverse each keyword in the column values provided
        for keyword in column_values:
            # check if the main findings data has any keywords mentioned in it
            new_df = main_excel_data[main_excel_data[main_column_name].str.contains(keyword)]
            # append all findings into a new dataframe
            appended_data.append(new_df)

        # concat all appended values into a single dataframe
        appended_data = pd.concat(appended_data)
        # remove all duplicate values in the dataframe
        appended_data = appended_data.drop_duplicates()
        # write the relevant findings data to the specified excel sheet
        appended_data.to_excel(file_path, index=False, engine='xlsxwriter')
        # write the recommendation for each header found in a doc.
        write_to_doc(heading, list_to_write, remed_filepath)

    return

def write_to_doc(heading, list_to_write, target_doc):
    """
    A function to write headline, data inside a document - .docx
    :param heading:
    :param list_to_write:
    :param target_doc:
    :return:
    """
    # create a docx object
    my_doc = docx.Document()
    my_doc.add_paragraph(heading, style= 'Title')

    try :
        for each_recommendation in list_to_write:
            # write the list of recommendations
            my_doc.add_paragraph(each_recommendation, style='List Bullet')

        # save the document created
        my_doc.save(target_doc)
        return

    except KeyError:
        return

def if_file_exists(file_name):
    """
    A function to check if any file exists in the specified path
    :param file_name:
    :return:
    """
    file_exist = os.path.exists(file_name)
    return (file_exist)

def func_bar(a, b):
    """
    A function to generate a loading bar for beautification purposes.
    :param a:
    :param b:
    :return NONE:
    """
    for i in tqdm(range(a), desc=b, ascii=False):
        time.sleep(0.01)
    return



